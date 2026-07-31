
from docx import Document
import os

class DocumentGenerator:
    def save(self, content):
        doc = Document()
        doc.add_heading("AI Generated Report", 1)
        doc.add_paragraph(content)

        # Create output directory if it doesn't exist
        output_dir = os.path.join(os.getcwd(), "output")
        os.makedirs(output_dir, exist_ok=True)

        path = os.path.join(output_dir, "res.docx")

        doc.save(path)
        return path
